"""Utilidades compartidas para generar documentos SENA usando la plantilla oficial GFPI-F-135."""
import shutil
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Paleta monocromática print-safe (regla dura del proyecto)
DARK = RGBColor(0x2C, 0x2C, 0x2C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_HEX = "2C2C2C"
MID_HEX = "D0D0D0"
LIGHT_HEX = "F0F0F0"

# Ruta a la plantilla oficial (ubicada en templates/)
TEMPLATE_PATH = Path(__file__).parent.parent / "templates" / "GFPI-F-135.docx"


def clone_template(destination_path: str) -> Document:
    """Copia la plantilla oficial y devuelve un Document listo para llenar,
    preservando encabezado, código GFPI-F-135, logo SENA y pie de página."""
    shutil.copy(TEMPLATE_PATH, destination_path)
    doc = Document(destination_path)

    body = doc.element.body
    # Vaciar el cuerpo, preservando el sectPr (que referencia headers/footers)
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    return doc


def set_run(run, bold=False, italic=False, size=11, color=DARK):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_heading(doc, text, size=13, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_run(r, bold=True, size=size)
    return p


def add_sub(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, bold=True, size=11)
    return p


def add_field(doc, label, value):
    """Renglón tipo 'Etiqueta: valor' con etiqueta en negrilla."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label + " ")
    set_run(r1, bold=True)
    r2 = p.add_run(str(value) if value else "")
    set_run(r2)
    return p


def add_text(doc, text, bold=False, italic=False, size=11, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_run(r, bold=bold, italic=italic, size=size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("•  " + text)
    set_run(r)
    return p


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def make_table(doc, rows, widths_cm, header_shading=DARK_HEX):
    """Crea una tabla con la primera fila como encabezado oscuro y filas alternas grises."""
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = False
    for j, w in enumerate(widths_cm):
        for row in table.rows:
            row.cells[j].width = Cm(w)
    for i, row_data in enumerate(rows):
        for j, text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(text) if text else "")
            if i == 0:
                set_run(r, bold=True, size=10, color=WHITE)
                set_cell_shading(cell, header_shading)
            else:
                set_run(r, size=10)
                if i % 2 == 0:
                    set_cell_shading(cell, LIGHT_HEX)
    _add_table_borders(table)
    return table


def _add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        borders.append(b)
    tblPr.append(borders)
